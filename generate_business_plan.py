"""
Génère un document Word professionnel contenant le plan de commercialisation
de BrightCalendar (marché québécois — lavage de vitres et services).
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


BRAND_BLUE = RGBColor(0x08, 0x91, 0xB2)
DARK = RGBColor(0x0A, 0x0A, 0x0A)
GREY = RGBColor(0x73, 0x73, 0x73)
GREEN = RGBColor(0x10, 0xB9, 0x81)
ORANGE = RGBColor(0xF5, 0x9E, 0x0B)


def shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def set_cell_border(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for border in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{border}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), 'CCCCCC')
        borders.append(b)
    tc_pr.append(borders)


def add_heading(doc, text, level=1, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 0:
        run.font.size = Pt(26)
    elif level == 1:
        run.font.size = Pt(18)
    elif level == 2:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)
    if color:
        run.font.color.rgb = color
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_para(doc, text, bold=False, italic=False, size=11, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(11)
        p.add_run(text).font.size = Pt(11)
    else:
        p.add_run(text).font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)
    return p


def build_table(doc, headers, rows, header_color='0891B2', col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Header
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(11)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade_cell(hdr[i], header_color)
        set_cell_border(hdr[i])
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for r_idx, row_data in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            cells[c_idx].text = ''
            p = cells[c_idx].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            set_cell_border(cells[c_idx])
            cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # Alternate row shading
            if r_idx % 2 == 0:
                shade_cell(cells[c_idx], 'F9FAFB')

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    return table


def add_spacer(doc, size=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(size)


# ============================================================
# DOCUMENT GENERATION
# ============================================================
doc = Document()

# Set default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Margins
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

# ============================================================
# TITLE PAGE
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('BrightCalendar')
r.bold = True
r.font.size = Pt(42)
r.font.color.rgb = BRAND_BLUE

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run('Plan de Commercialisation & Analyse Financière')
r.font.size = Pt(16)
r.italic = True
r.font.color.rgb = GREY

add_spacer(doc, 24)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = info.add_run('Application SaaS de gestion de rendez-vous\n')
r.font.size = Pt(12)
r = info.add_run('Spécialisée pour les entreprises de lavage de vitres\net autres services à domicile')
r.font.size = Pt(12)
r.italic = True

add_spacer(doc, 36)

author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = author.add_run('Préparé pour: Lavage de Vitres Bois-Franc\n')
r.font.size = Pt(11)
r.font.color.rgb = DARK
r = author.add_run('Document de référence — Stratégie commerciale\n')
r.font.size = Pt(11)
r.italic = True
r.font.color.rgb = GREY
r = author.add_run('Version 1.0 — Juin 2025')
r.font.size = Pt(10)
r.font.color.rgb = GREY

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (manual)
# ============================================================
add_heading(doc, 'Table des matières', level=1, color=BRAND_BLUE)
toc = [
    '1. Résumé exécutif',
    '2. Architecture commerciale proposée',
    '3. Coûts de développement (crédits Emergent)',
    '4. Coûts de publication — App Store Apple',
    '5. Coûts de publication — Google Play',
    '6. Modèles de tarification',
    '7. Simulation de rentabilité',
    '8. Feuille de route recommandée',
    '9. Prochaines étapes',
]
for item in toc:
    add_para(doc, item, size=11)

doc.add_page_break()

# ============================================================
# 1. EXECUTIVE SUMMARY
# ============================================================
add_heading(doc, '1. Résumé exécutif', level=1, color=BRAND_BLUE)

add_para(doc,
    "BrightCalendar est une application mobile et web complète, déjà développée et fonctionnelle, "
    "destinée aux entreprises de services à domicile (lavage de vitres, entretien, plomberie, etc.). "
    "Elle intègre la gestion de rendez-vous multi-vues, une base de données clients (CRM), "
    "la génération de factures PDF, les campagnes courriel saisonnières, le formulaire public de réservation, "
    "la dictée vocale (Whisper IA), un QR Code personnalisé et une interface trilingue (FR/EN/ES).",
    size=11)

add_heading(doc, 'Opportunité de marché', level=2, color=DARK)
add_bullet(doc, "Au Québec seulement : +5 000 entreprises de lavage de vitres, entretien ménager et services extérieurs.", "🎯 ")
add_bullet(doc, "La plupart utilisent encore Excel, papier ou des agendas génériques non spécialisés.", "📊 ")
add_bullet(doc, "Aucune solution francophone dédiée à ce secteur sur le marché québécois.", "🇨🇦 ")
add_bullet(doc, "Potentiel d'expansion : provinces canadiennes anglophones, USA, Amérique latine (app trilingue).", "🌍 ")

add_heading(doc, 'Avantages compétitifs', level=2, color=DARK)
add_bullet(doc, "Application déjà codée — pas de cycle de développement long à anticiper.", "✅ ")
add_bullet(doc, "Interface native mobile (iOS + Android) + version web.", "✅ ")
add_bullet(doc, "Support trilingue (FR/EN/ES) déjà intégré — marché élargi.", "✅ ")
add_bullet(doc, "Architecture modulaire : certaines fonctionnalités peuvent être activées ou désactivées.", "✅ ")
add_bullet(doc, "Marge brute potentielle très élevée (>90 %) une fois le développement terminé.", "✅ ")

doc.add_page_break()

# ============================================================
# 2. ARCHITECTURE
# ============================================================
add_heading(doc, '2. Architecture commerciale proposée', level=1, color=BRAND_BLUE)

add_para(doc,
    "Pour rendre BrightCalendar commercialisable à plusieurs entreprises, l'application doit évoluer "
    "selon trois couches techniques, que nous recommandons de déployer progressivement :", size=11)

add_heading(doc, 'Couche 1 — Multi-Tenant (isolation des données)', level=2, color=GREEN)
add_bullet(doc, "Chaque entreprise cliente dispose de son propre compte isolé (company_id).")
add_bullet(doc, "Authentification par courriel/mot de passe (JWT).")
add_bullet(doc, "Toutes les requêtes API filtrent automatiquement par company_id.")
add_bullet(doc, "Aucune fuite de données entre entreprises — conformité RGPD / Loi 25 (Québec).")

add_heading(doc, 'Couche 2 — Paramètres Entreprise', level=2, color=GREEN)
add_para(doc,
    "Une page /settings permettant à chaque propriétaire de configurer sa version de l'app :",
    size=11)

headers = ['Paramètre', 'Exemple / Usage']
rows = [
    ['🏢 Identité', 'Nom, logo, couleur de marque, slogan'],
    ['📞 Coordonnées', 'Téléphone, courriel, site web, adresse'],
    ['💰 Prix estimation', 'Prix par fenêtre, porte, puits de lumière'],
    ['🧾 Taxes', 'Numéros TPS/TVQ, texte de pied de facture'],
    ['✉️ Modèles courriels', 'Sujets et corps des 3 campagnes saisonnières'],
    ['🗓️ Horaires', 'Jours et heures de travail par défaut'],
    ['🌍 Langue défaut', 'FR, EN ou ES'],
]
build_table(doc, headers, rows, col_widths=[4.5, 11])

add_heading(doc, 'Couche 3 — Feature Flags (modules activables)', level=2, color=GREEN)
add_para(doc,
    "Chaque module peut être activé ou désactivé selon le forfait acheté par l'entreprise cliente :",
    size=11)
modules = [
    'Campagnes Marketing courriel',
    'Base Clients (CRM) + Import Excel/CSV',
    'VIP List (formulaire public)',
    'Employés multiples (désactivé pour entreprise solo)',
    'Statistiques avancées',
    'Estimation en ligne publique',
    'QR Code de réservation',
    'Dictée vocale (IA Whisper)',
    'Synchronisation Google Calendar',
    'Rappels SMS automatiques',
]
for m in modules:
    add_bullet(doc, m, "☑ ")

doc.add_page_break()

# ============================================================
# 3. COÛTS DÉVELOPPEMENT
# ============================================================
add_heading(doc, '3. Coûts de développement (crédits Emergent)', level=1, color=BRAND_BLUE)
add_para(doc,
    "Estimation indicative pour compléter la version commercialisable. "
    "Les valeurs peuvent varier ±30 % selon la complexité rencontrée.", italic=True, color=GREY, size=10)

headers = ['Phase / Fonctionnalité', 'Description', 'Crédits estimés']
rows = [
    ['Phase 1 — Paramètres entreprise',
     'Logo upload, coordonnées, prix estimation, taxes, intégration factures',
     '15 à 25'],
    ['Phase 2 — Paiements Stripe',
     'Modules activables + lien de paiement en ligne sur factures PDF',
     '25 à 40'],
    ['Phase 3 — Multi-tenant SaaS',
     'Authentification, company_id, migrations DB, isolation complète',
     '60 à 100'],
    ['Push Notifications + SMS',
     'Expo Push Notifications + rappels Twilio 24 h avant RDV',
     '20 à 30'],
    ['Google Calendar Sync',
     'OAuth Google + synchronisation bidirectionnelle RDV',
     '25 à 40'],
    ['Débogage, tests, polish UX',
     'Ajustements finaux, correction de bogues, optimisation',
     '20 à 30'],
    ['TOTAL version complète',
     'Toutes les phases cumulées',
     '165 à 265 crédits'],
]
build_table(doc, headers, rows, col_widths=[5, 8.5, 3])

add_spacer(doc, 6)
add_para(doc, "💡 Version minimum vendable (Phase 1 + Stripe) : environ 40 à 65 crédits.",
         bold=True, color=ORANGE, size=11)

doc.add_page_break()

# ============================================================
# 4. APPLE APP STORE
# ============================================================
add_heading(doc, '4. Coûts de publication — App Store Apple', level=1, color=BRAND_BLUE)

headers = ['Item', 'Coût', 'Fréquence']
rows = [
    ['Apple Developer Program', '99 USD (~135 CAD)', 'Annuel — obligatoire'],
    ['Apple Developer Enterprise', '299 USD (~405 CAD)', 'Annuel (optionnel, distribution interne seulement)'],
    ['Commission sur ventes in-app', '30 % (15 % si < 1 M $ / an)', 'Par transaction'],
    ['Build iOS via Expo EAS', 'Gratuit jusqu\'à 30 builds/mois', '19 à 99 USD/mois ensuite'],
    ['Délai de révision', '24 à 48 h en moyenne', 'À chaque soumission'],
]
build_table(doc, headers, rows, header_color='000000', col_widths=[5.5, 5.5, 5.5])

add_spacer(doc)
add_heading(doc, '⚠️ Avertissements importants Apple', level=2, color=ORANGE)
add_bullet(doc, "Apple exige souvent une authentification utilisateur pour les apps SaaS B2B — la Phase 3 (multi-tenant) sera probablement requise avant publication.")
add_bullet(doc, "Les politiques Apple interdisent de rediriger vers des paiements externes — Stripe doit être implémenté correctement (paiements in-app ou webview conforme).")
add_bullet(doc, "L'app doit être utilisable immédiatement après installation (pas de page blanche exigeant un abonnement avant la première interaction).")

doc.add_page_break()

# ============================================================
# 5. GOOGLE PLAY
# ============================================================
add_heading(doc, '5. Coûts de publication — Google Play Store', level=1, color=BRAND_BLUE)

headers = ['Item', 'Coût', 'Fréquence']
rows = [
    ['Google Play Developer Account', '25 USD (~34 CAD)', 'Frais unique à vie 🎉'],
    ['Commission sur ventes in-app', '30 % (15 % si < 1 M $ / an)', 'Par transaction'],
    ['Build Android via Expo EAS', 'Gratuit jusqu\'à 30 builds/mois', '19 à 99 USD/mois ensuite'],
    ['Délai de révision', '1 à 7 jours', 'À chaque soumission'],
]
build_table(doc, headers, rows, header_color='34A853', col_widths=[5.5, 5.5, 5.5])

add_spacer(doc)
add_para(doc,
    "✅ Google Play est environ 10 fois moins cher qu'Apple et beaucoup plus permissif. "
    "Nous recommandons de publier d'abord sur Android pour valider le marché, puis iOS en second.",
    bold=True, color=GREEN, size=11)

doc.add_page_break()

# ============================================================
# 6. MODÈLES DE TARIFICATION
# ============================================================
add_heading(doc, '6. Modèles de tarification proposés', level=1, color=BRAND_BLUE)

# Option A
add_heading(doc, 'Option A — Licence One-Shot', level=2, color=DARK)
add_para(doc, "Vous livrez une version pré-configurée à chaque entreprise. Simple à démarrer, marge très élevée.", size=11)
headers = ['Forfait', 'Prix CAD', 'Inclus']
rows = [
    ['🥉 Starter', '499 $', 'Calendrier, CRM, factures PDF'],
    ['🥈 Pro', '999 $', '+ Campagnes marketing, VIP List, estimation en ligne'],
    ['🥇 Business', '1 499 $', '+ QR, stats, dictée vocale, multi-employés'],
    ['🎨 Setup personnalisé', '+ 250 $', 'Upload logo, prix, coordonnées (add-on)'],
]
build_table(doc, headers, rows, col_widths=[4, 3, 9])
add_spacer(doc)

# Option B
add_heading(doc, 'Option B — Abonnement SaaS (revenus récurrents)', level=2, color=DARK)
add_para(doc, "Nécessite la Phase 3 (multi-tenant). Génère un revenu mensuel récurrent (MRR) stable.", size=11)
headers = ['Forfait', 'Mensuel', 'Annuel (–20 %)', 'Cible']
rows = [
    ['🆓 Free', '0 $', '0 $', '1 utilisateur, 20 clients, logo Emergent visible'],
    ['🥉 Starter', '29 $/mois', '290 $/an', 'Artisan solo, lavage de vitres'],
    ['🥈 Pro', '59 $/mois', '590 $/an', 'PME 2 à 5 employés'],
    ['🥇 Business', '99 $/mois', '990 $/an', 'PME 5 à 15 employés, multi-sites'],
]
build_table(doc, headers, rows, col_widths=[3, 3, 3, 6.5])
add_spacer(doc)

# Option C
add_heading(doc, 'Option C — Hybride (recommandée)', level=2, color=GREEN)
add_para(doc,
    "Combine les avantages des deux modèles précédents. Client paie moins cher à l'achat, "
    "mais vous conservez un revenu récurrent pour le support et les mises à jour.",
    size=11)
headers = ['Offre', 'Prix', 'Description']
rows = [
    ['🎁 Installation personnalisée', '499 $ (one-shot)', 'Setup, branding, formation de base'],
    ['🔁 Support + Mises à jour', '19 $/mois ou 190 $/an', 'Nouvelles fonctionnalités, correctifs, assistance'],
]
build_table(doc, headers, rows, col_widths=[5, 5, 6])

doc.add_page_break()

# ============================================================
# 7. SIMULATION RENTABILITÉ
# ============================================================
add_heading(doc, '7. Simulation de rentabilité', level=1, color=BRAND_BLUE)
add_para(doc,
    "Scénario réaliste avec l'Option C (hybride) : vous vendez à 5 entreprises québécoises de services à domicile la première année.",
    italic=True, size=11)

add_heading(doc, 'Revenus Année 1', level=2, color=GREEN)
headers = ['Source', 'Détail', 'Revenus']
rows = [
    ['Installations', '5 × 499 $', '2 495 $'],
    ['Abonnements annuels', '5 × 190 $/an', '950 $'],
    ['TOTAL Année 1', '', '3 445 $ CAD'],
]
build_table(doc, headers, rows, col_widths=[5, 5, 5])
add_spacer(doc)

add_heading(doc, 'Revenus Année 2+ (cumulatif)', level=2, color=GREEN)
headers = ['Année', 'Clients cumulés', 'Revenus abonnement', 'Revenus installation']
rows = [
    ['Année 1', '5', '950 $', '2 495 $'],
    ['Année 2', '12', '2 280 $', '3 493 $'],
    ['Année 3', '25', '4 750 $', '6 487 $'],
    ['Année 5', '60', '11 400 $', '15 000 $+'],
]
build_table(doc, headers, rows, col_widths=[3, 3.5, 4.5, 4.5])
add_spacer(doc)

add_heading(doc, 'Coûts fixes annuels', level=2, color=ORANGE)
headers = ['Item', 'Coût annuel CAD']
rows = [
    ['Apple Developer Program', '135 $'],
    ['Google Play (one-shot — amortie sur 5 ans)', '7 $'],
    ['Hébergement (inclus Emergent)', '0 $'],
    ['Crédits Emergent (maintenance ~30 crédits/an)', 'Variable'],
    ['TOTAL coûts fixes', '~ 140 à 300 $/an'],
]
build_table(doc, headers, rows, col_widths=[9, 5])
add_spacer(doc)

p = doc.add_paragraph()
r = p.add_run("➡️ Marge nette projetée Année 1 : environ 3 100 $ CAD pour 5 ventes.")
r.bold = True
r.font.color.rgb = GREEN
r.font.size = Pt(13)

p = doc.add_paragraph()
r = p.add_run("➡️ Marge nette projetée Année 3 : environ 10 500 $ CAD (pour 25 clients actifs).")
r.bold = True
r.font.color.rgb = GREEN
r.font.size = Pt(13)

doc.add_page_break()

# ============================================================
# 8. FEUILLE DE ROUTE
# ============================================================
add_heading(doc, '8. Feuille de route recommandée', level=1, color=BRAND_BLUE)

steps = [
    ("Étape 1 — Phase 1 (Paramètres Entreprise)",
     "1 à 2 sessions",
     "Rendre l'app vendable rapidement. Logo, coordonnées, prix estimation modifiables."),
    ("Étape 2 — Intégration Stripe",
     "1 session",
     "Accepter les paiements en ligne pour factures clients et abonnements."),
    ("Étape 3 — Publication Google Play",
     "1 à 2 jours",
     "Frais unique 25 USD. Approbation rapide. Commencer à vendre."),
    ("Étape 4 — Premières ventes (3-5 clients)",
     "1 à 3 mois",
     "Vente Option A (licence one-shot) pour valider le marché."),
    ("Étape 5 — Phase 2 (Feature Flags)",
     "1 à 2 sessions",
     "Créer les forfaits Starter/Pro/Business avec modules activables."),
    ("Étape 6 — Publication Apple App Store",
     "1 semaine",
     "Frais 99 USD/an. Révision plus stricte. Cible clientèle iOS."),
    ("Étape 7 — Phase 3 (Multi-tenant SaaS)",
     "3 à 5 sessions",
     "Auth, isolation DB. Transition vers modèle d'abonnement récurrent (MRR)."),
    ("Étape 8 — Marketing & expansion",
     "En continu",
     "Site web, SEO, publicité Facebook/Google, témoignages clients."),
]

for i, (title_step, duration, desc) in enumerate(steps, 1):
    p = doc.add_paragraph()
    r = p.add_run(f"{title_step}  ")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = BRAND_BLUE
    r = p.add_run(f"({duration})")
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = GREY
    add_para(doc, f"   {desc}", size=11)
    add_spacer(doc, 3)

doc.add_page_break()

# ============================================================
# 9. PROCHAINES ÉTAPES
# ============================================================
add_heading(doc, '9. Prochaines étapes immédiates', level=1, color=BRAND_BLUE)

add_para(doc,
    "Actions concrètes à entreprendre dans l'ordre suggéré :",
    size=11)

add_bullet(doc, "Valider cette stratégie et choisir le modèle de tarification (A, B ou C).", "✅ ")
add_bullet(doc, "Démarrer le développement de la Phase 1 (Paramètres Entreprise).", "🛠️ ")
add_bullet(doc, "Créer un compte Google Play Developer (25 USD — à vie).", "🤖 ")
add_bullet(doc, "Préparer 2-3 démonstrations pour prospects québécois (lavage de vitres).", "🎬 ")
add_bullet(doc, "Créer un site vitrine simple avec formulaire de contact.", "🌐 ")
add_bullet(doc, "Signer le premier contrat client (installation one-shot 499 $).", "💰 ")
add_bullet(doc, "Intégrer Stripe et préparer la Phase 2 (Feature Flags).", "💳 ")
add_bullet(doc, "À 5-10 clients satisfaits : démarrer la Phase 3 (multi-tenant SaaS).", "🚀 ")

add_spacer(doc, 18)

note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = note.add_run("Document généré automatiquement — BrightCalendar © 2025\n")
r.font.size = Pt(9)
r.font.color.rgb = GREY
r = note.add_run("Lavage de Vitres Bois-Franc — Stratégie de commercialisation")
r.font.size = Pt(9)
r.italic = True
r.font.color.rgb = GREY

# ============================================================
# SAVE
# ============================================================
output_path = '/app/backend/assets/BrightCalendar_Plan_Commercialisation.docx'
doc.save(output_path)
print(f"Document généré avec succès: {output_path}")
