"""
Génère une BROCHURE COMMERCIALE professionnelle pour BrightCalendar.
Destinée aux prospects clients (entreprises de services à domicile).
Aucune mention de crédits Emergent ou de coûts internes.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


BRAND_BLUE = RGBColor(0x08, 0x91, 0xB2)
DARK = RGBColor(0x0A, 0x0A, 0x0A)
GREY = RGBColor(0x73, 0x73, 0x73)
GREEN = RGBColor(0x10, 0xB9, 0x81)
GOLD = RGBColor(0xF5, 0x9E, 0x0B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def set_cell_border(cell, color='CCCCCC'):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for border in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{border}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), color)
        borders.append(b)
    tc_pr.append(borders)


def add_heading(doc, text, level=1, color=None, center=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    sizes = {0: 36, 1: 22, 2: 16, 3: 13}
    run.font.size = Pt(sizes.get(level, 11))
    if color:
        run.font.color.rgb = color
    if center or level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_para(doc, text, bold=False, italic=False, size=11, color=None, center=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc, text, bold_prefix=None, size=11):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(size)
        p.add_run(text).font.size = Pt(size)
    else:
        p.add_run(text).font.size = Pt(size)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_feature_box(doc, icon, title, desc):
    """Adds a nicely formatted feature box with icon, title and description."""
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    cell_icon = table.rows[0].cells[0]
    cell_text = table.rows[0].cells[1]
    cell_icon.width = Cm(1.8)
    cell_text.width = Cm(13.5)

    # Icon cell
    cell_icon.text = ''
    p = cell_icon.paragraphs[0]
    r = p.add_run(icon)
    r.font.size = Pt(22)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_icon.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Text cell
    cell_text.text = ''
    p = cell_text.paragraphs[0]
    r = p.add_run(title + '\n')
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = BRAND_BLUE
    r = p.add_run(desc)
    r.font.size = Pt(10)
    r.font.color.rgb = DARK

    add_para(doc, '', size=4)  # small spacer


def build_table(doc, headers, rows, header_color='0891B2', col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(11)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade_cell(hdr[i], header_color)
        set_cell_border(hdr[i])
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r_idx, row_data in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            cells[c_idx].text = ''
            p = cells[c_idx].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            set_cell_border(cells[c_idx])
            cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_idx % 2 == 0:
                shade_cell(cells[c_idx], 'F9FAFB')

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


def add_spacer(doc, size=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(size)


def add_hr(doc):
    """Adds a horizontal divider line."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━')
    r.font.color.rgb = BRAND_BLUE
    r.font.size = Pt(10)


# ============================================================
# DOCUMENT
# ============================================================
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

# ============================================================
# COUVERTURE
# ============================================================
add_spacer(doc, 40)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('BrightCalendar')
r.bold = True
r.font.size = Pt(54)
r.font.color.rgb = BRAND_BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('✨')
r.font.size = Pt(36)

add_spacer(doc, 24)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("L'application tout-en-un pour\ngérer votre entreprise de services")
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = DARK

add_spacer(doc, 18)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Calendrier • Clients • Facturation • Marketing\nEstimation en ligne • Campagnes courriel")
r.italic = True
r.font.size = Pt(14)
r.font.color.rgb = GREY

add_spacer(doc, 60)

# Tagline banner
tag_table = doc.add_table(rows=1, cols=1)
tc = tag_table.rows[0].cells[0]
shade_cell(tc, '0891B2')
tc.text = ''
p = tc.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("💼 Prêt à l'emploi  •  📱 Mobile & Web  •  🇨🇦 Conçu au Québec")
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = WHITE

add_spacer(doc, 30)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Brochure commerciale 2025")
r.font.size = Pt(11)
r.italic = True
r.font.color.rgb = GREY

doc.add_page_break()

# ============================================================
# PAGE 2 — POURQUOI
# ============================================================
add_heading(doc, "Pourquoi BrightCalendar?", level=1, color=BRAND_BLUE, center=True)
add_hr(doc)
add_spacer(doc, 10)

add_para(doc,
    "Vous êtes propriétaire d'une entreprise de services à domicile — lavage de vitres, entretien, "
    "plomberie, électricité, ménage, paysagement? Vous jonglez avec un agenda papier, des fichiers Excel, "
    "des textos Facebook Messenger et des factures Word?",
    size=12)

add_spacer(doc, 6)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Il est temps de passer au 21e siècle. 🚀")
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = GREEN

add_spacer(doc, 12)

add_heading(doc, "Ce que BrightCalendar vous apporte :", level=2, color=DARK)

add_feature_box(doc, "📅",
    "Agenda intelligent multi-vues",
    "Visualisez vos rendez-vous en jour, semaine ou mois. Glissez-déposez pour reporter. "
    "Fini les erreurs de double-réservation.")

add_feature_box(doc, "👥",
    "Base de données clients (CRM)",
    "Importez votre liste Excel en 10 secondes. Historique complet de chaque client, tags, notes, "
    "recherche instantanée.")

add_feature_box(doc, "🧾",
    "Factures PDF professionnelles",
    "Générez une facture en 1 clic avec votre logo, votre numéro de TPS/TVQ et vos coordonnées. "
    "Envoi direct par courriel au client.")

add_feature_box(doc, "💰",
    "Outil d'estimation intégré",
    "Calculez un prix en temps réel selon le nombre de fenêtres, portes et puits de lumière. "
    "Partagez l'estimation avec votre client.")

add_feature_box(doc, "📧",
    "Campagnes marketing saisonnières",
    "3 campagnes pré-rédigées (Printemps, Automne, Été) pour relancer vos clients aux bons moments "
    "de l'année. Personnalisables en 3 langues.")

add_feature_box(doc, "🌐",
    "Formulaire de réservation public",
    "Vos clients peuvent prendre rendez-vous directement en ligne, 24 h sur 24. QR Code intégré "
    "pour vos cartes d'affaires et flyers.")

add_feature_box(doc, "🎤",
    "Dictée vocale intelligente",
    "Dictez vos notes de rendez-vous — l'intelligence artificielle les transcrit automatiquement. "
    "Gagnez du temps sur la route.")

add_feature_box(doc, "🌍",
    "Trilingue (FR / EN / ES)",
    "Votre app change automatiquement de langue selon votre préférence. Parfait si vous servez "
    "une clientèle diversifiée.")

doc.add_page_break()

# ============================================================
# PAGE 3 — COMPARATIF
# ============================================================
add_heading(doc, "Avant / Après BrightCalendar", level=1, color=BRAND_BLUE, center=True)
add_hr(doc)
add_spacer(doc, 10)

headers = ['Situation', '❌ Avant', '✅ Avec BrightCalendar']
rows = [
    ['Prise de rendez-vous', 'Textos, appels manqués, agenda papier', 'Formulaire en ligne 24/7'],
    ['Liste de clients', 'Fichier Excel perdu, feuilles volantes', 'CRM cloud synchronisé partout'],
    ['Factures', 'Word/Excel, envoi manuel, erreurs fréquentes', 'PDF généré en 1 clic, logo inclus'],
    ['Estimations', 'Calcul mental, erreurs de prix', 'Calcul automatique + partage client'],
    ['Campagnes courriel', 'Aucune ou sporadiques', '3 campagnes saisonnières automatisées'],
    ['Notes de RDV', 'Calepin, notes sur le téléphone', 'Dictée vocale + archivage automatique'],
    ['Rappels clients', 'Oubliés, perdus', 'Notification automatique 24h avant'],
    ['Suivi du revenu', 'Fin d\'année stressante', 'Statistiques en temps réel'],
]
build_table(doc, headers, rows, col_widths=[4.5, 6.5, 5.5])

add_spacer(doc, 14)

add_heading(doc, "Gain de temps estimé :", level=2, color=GREEN, center=True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("⏱️  5 à 10 heures par semaine")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = GREEN

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Soit 20 à 40 heures par mois pour développer votre entreprise, voir vos clients — ou votre famille.")
r.italic = True
r.font.size = Pt(11)
r.font.color.rgb = GREY

doc.add_page_break()

# ============================================================
# PAGE 4 — FORFAITS
# ============================================================
add_heading(doc, "Nos forfaits", level=1, color=BRAND_BLUE, center=True)
add_hr(doc)
add_spacer(doc, 10)

add_para(doc,
    "Trois formules pensées pour chaque taille d'entreprise. Tous les forfaits incluent l'hébergement "
    "cloud sécurisé, les mises à jour gratuites à vie et le support courriel.",
    italic=True, size=11, center=True)

add_spacer(doc, 10)

# ---- STARTER ----
add_heading(doc, "🥉 Starter — L'essentiel", level=2, color=DARK)
p = doc.add_paragraph()
r = p.add_run("499 $ CAD")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = GREEN
p.add_run("  ")
r = p.add_run("(licence unique)")
r.italic = True
r.font.size = Pt(11)
r.font.color.rgb = GREY

add_bullet(doc, "Calendrier multi-vues (jour / semaine / mois)", "✓ ")
add_bullet(doc, "Base de clients illimitée + import Excel/CSV", "✓ ")
add_bullet(doc, "Factures PDF avec votre logo", "✓ ")
add_bullet(doc, "Dictée vocale pour notes de RDV", "✓ ")
add_bullet(doc, "Support courriel inclus", "✓ ")
add_para(doc, "👤 Idéal pour : artisan solo ou entreprise de 1-2 personnes", italic=True, color=GREY, size=10)

add_spacer(doc, 14)

# ---- PRO ----
add_heading(doc, "🥈 Pro — Le plus populaire ⭐", level=2, color=BRAND_BLUE)
p = doc.add_paragraph()
r = p.add_run("999 $ CAD")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = GREEN
p.add_run("  ")
r = p.add_run("(licence unique)")
r.italic = True
r.font.size = Pt(11)
r.font.color.rgb = GREY

add_bullet(doc, "Tout ce qui est inclus dans Starter", "✓ ")
add_bullet(doc, "Campagnes marketing courriel (3 saisons)", "✓ ")
add_bullet(doc, "Liste VIP (formulaire public d'inscription)", "✓ ")
add_bullet(doc, "Outil d'estimation en ligne partageable", "✓ ")
add_bullet(doc, "QR Code de réservation personnalisé", "✓ ")
add_bullet(doc, "Mode trilingue FR / EN / ES", "✓ ")
add_para(doc, "👥 Idéal pour : PME de 2 à 5 employés", italic=True, color=GREY, size=10)

add_spacer(doc, 14)

# ---- BUSINESS ----
add_heading(doc, "🥇 Business — Pour les plus ambitieux", level=2, color=GOLD)
p = doc.add_paragraph()
r = p.add_run("1 499 $ CAD")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = GREEN
p.add_run("  ")
r = p.add_run("(licence unique)")
r.italic = True
r.font.size = Pt(11)
r.font.color.rgb = GREY

add_bullet(doc, "Tout ce qui est inclus dans Pro", "✓ ")
add_bullet(doc, "Gestion multi-employés avec couleurs", "✓ ")
add_bullet(doc, "Tableau de bord statistiques avancées", "✓ ")
add_bullet(doc, "Historique complet des clients & revenus", "✓ ")
add_bullet(doc, "Sauvegarde quotidienne automatique", "✓ ")
add_bullet(doc, "Support téléphonique prioritaire", "✓ ")
add_para(doc, "🏢 Idéal pour : PME de 5 à 15 employés", italic=True, color=GREY, size=10)

add_spacer(doc, 14)

# Add-on
add_heading(doc, "🎨 Service de personnalisation (en option)", level=3, color=DARK)
p = doc.add_paragraph()
r = p.add_run("+ 250 $ CAD")
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = BRAND_BLUE
add_para(doc,
    "Nous configurons l'application à votre image : upload de votre logo, intégration de vos prix "
    "d'estimation, vos coordonnées, vos numéros de taxes. Vous recevez une application clé en main, "
    "prête à être utilisée dès le premier jour.",
    size=11)

doc.add_page_break()

# ============================================================
# PAGE 5 — ABONNEMENT SUPPORT
# ============================================================
add_heading(doc, "Option : Support + Mises à jour continues", level=1, color=BRAND_BLUE, center=True)
add_hr(doc)
add_spacer(doc, 10)

add_para(doc,
    "Pour un tarif symbolique, bénéficiez de l'évolution continue de BrightCalendar et d'une assistance "
    "réactive en cas de question.",
    size=12)

add_spacer(doc, 8)

headers = ['Formule', 'Prix', 'Inclus']
rows = [
    ['🔁 Mensuel', '19 $ / mois', 'Mises à jour illimitées + support courriel sous 24 h'],
    ['💎 Annuel (économie 20 %)', '190 $ / an', 'Tout ce qui précède + priorité de support + nouvelles fonctionnalités'],
]
build_table(doc, headers, rows, col_widths=[4.5, 3.5, 8.5])

add_spacer(doc, 12)

add_heading(doc, "Fonctionnalités à venir (incluses dans l'abonnement)", level=2, color=DARK)
add_bullet(doc, "💳 Paiement en ligne sur factures (Stripe)", "🔜 ")
add_bullet(doc, "📆 Synchronisation Google Calendar", "🔜 ")
add_bullet(doc, "🔔 Rappels SMS automatiques 24 h avant RDV", "🔜 ")
add_bullet(doc, "📲 Notifications push pour nouvelles demandes", "🔜 ")
add_bullet(doc, "🌟 Gestion des avis clients (Google Reviews)", "🔜 ")

doc.add_page_break()

# ============================================================
# PAGE 6 — COMMENT ÇA MARCHE
# ============================================================
add_heading(doc, "Comment ça fonctionne?", level=1, color=BRAND_BLUE, center=True)
add_hr(doc)
add_spacer(doc, 10)

steps = [
    ("1️⃣", "Démonstration gratuite (30 min)",
     "Nous vous montrons l'application en direct, adaptée à votre secteur. Aucun engagement."),
    ("2️⃣", "Choix du forfait",
     "Starter, Pro ou Business — selon la taille de votre équipe et vos besoins."),
    ("3️⃣", "Configuration personnalisée (option)",
     "Nous installons votre logo, vos prix et vos coordonnées. Livraison en 48 h."),
    ("4️⃣", "Formation + activation",
     "Session de formation de 1 h (visio ou sur place). Vous êtes autonome immédiatement."),
    ("5️⃣", "Utilisation au quotidien",
     "Votre équipe utilise BrightCalendar sur iPhone, Android et ordinateur. Partout, tout le temps."),
    ("6️⃣", "Support continu (option)",
     "Si vous avez choisi l'abonnement support, vous bénéficiez des mises à jour et de notre assistance."),
]
for num, title_s, desc in steps:
    p = doc.add_paragraph()
    r = p.add_run(num + "  ")
    r.font.size = Pt(18)
    r = p.add_run(title_s)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = BRAND_BLUE
    add_para(doc, "     " + desc, size=11)
    add_spacer(doc, 4)

add_spacer(doc, 12)

add_heading(doc, "Garantie satisfaction 30 jours", level=2, color=GREEN, center=True)
add_para(doc,
    "Si BrightCalendar ne vous apporte pas la valeur attendue dans les 30 premiers jours, "
    "nous vous remboursons intégralement. Sans question. Sans discussion.",
    italic=True, size=12, center=True)

doc.add_page_break()

# ============================================================
# PAGE 7 — FAQ
# ============================================================
add_heading(doc, "Questions fréquentes", level=1, color=BRAND_BLUE, center=True)
add_hr(doc)
add_spacer(doc, 10)

faqs = [
    ("🔒 Mes données sont-elles en sécurité?",
     "Oui. Vos données sont stockées dans un centre de données au Canada, chiffrées en transit et au repos. "
     "Sauvegarde automatique quotidienne. Conformité avec la Loi 25 du Québec."),
    ("📱 Sur quels appareils fonctionne BrightCalendar?",
     "iPhone, iPad, Android (téléphone et tablette), ainsi que tout ordinateur (Mac, PC) via un navigateur web. "
     "Les données sont synchronisées instantanément entre tous vos appareils."),
    ("🌐 Faut-il une connexion Internet permanente?",
     "Non. L'application fonctionne en mode hors ligne pour la consultation. Les modifications se "
     "synchronisent dès la reconnexion."),
    ("👥 Combien de clients puis-je ajouter?",
     "Aucune limite. Vous pouvez importer 10, 100 ou 10 000 clients sans frais supplémentaires."),
    ("🔄 Puis-je changer de forfait plus tard?",
     "Absolument. Vous pouvez monter en forfait (Starter → Pro → Business) à tout moment en payant "
     "seulement la différence."),
    ("💳 Quels sont les modes de paiement acceptés?",
     "Virement Interac, carte de crédit (Visa, Mastercard, AMEX), ou chèque d'entreprise. Facture avec "
     "TPS/TVQ incluse."),
    ("📧 Le support est-il en français?",
     "Oui, 100 %. Nous sommes basés au Québec et notre équipe parle français (et anglais au besoin)."),
    ("🆓 Y a-t-il un essai gratuit?",
     "Nous offrons une démonstration en direct de 30 minutes gratuitement, sans engagement. "
     "Vous pouvez aussi demander un accès test de 7 jours sur demande."),
]
for q, a in faqs:
    p = doc.add_paragraph()
    r = p.add_run(q)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = BRAND_BLUE
    add_para(doc, a, size=11)
    add_spacer(doc, 4)

doc.add_page_break()

# ============================================================
# PAGE 8 — CONTACT
# ============================================================
add_spacer(doc, 20)
add_heading(doc, "Prêt à moderniser votre entreprise?", level=1, color=BRAND_BLUE, center=True)
add_hr(doc)
add_spacer(doc, 14)

add_para(doc,
    "Réservez votre démonstration gratuite dès aujourd'hui. 30 minutes pour découvrir comment "
    "BrightCalendar peut transformer votre quotidien et faire gagner votre entreprise.",
    size=13, center=True)

add_spacer(doc, 24)

# Contact card
ct = doc.add_table(rows=1, cols=1)
tc = ct.rows[0].cells[0]
shade_cell(tc, '0891B2')
tc.text = ''
p = tc.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("\nCONTACT\n")
r.bold = True
r.font.color.rgb = WHITE
r.font.size = Pt(14)

r = p.add_run("\n📞  514-570-9802\n")
r.bold = True
r.font.color.rgb = WHITE
r.font.size = Pt(16)

r = p.add_run("🌐  https://Lavagedevitre.org\n")
r.bold = True
r.font.color.rgb = WHITE
r.font.size = Pt(13)

r = p.add_run("📧  contact@lavagedevitre.org\n\n")
r.bold = True
r.font.color.rgb = WHITE
r.font.size = Pt(13)

add_spacer(doc, 24)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("🎁  Mention « BROCHURE2025 » à la prise de contact :")
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = DARK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("100 $ de rabais sur l'installation personnalisée !")
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = GREEN

add_spacer(doc, 40)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("BrightCalendar — L'application pensée pour vous, propriétaire d'entreprise.\n")
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = GREY
r = p.add_run("Fait au Québec 🇨🇦 — © 2025 Tous droits réservés")
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = GREY

# SAVE
output_path = '/app/backend/assets/BrightCalendar_Brochure_Commerciale.docx'
doc.save(output_path)
print(f"Brochure générée avec succès: {output_path}")
